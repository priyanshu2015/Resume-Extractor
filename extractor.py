import docx2txt
import nltk
import re
from docx import Document
import docx
import PyPDF2
from pdfReader2 import convert_pdf_to_txt
import pandas as pd
import tabula
from pdfFontFinder import font_style
from pdfImageFinder import image_finder
from pdfPageIteratorusingFitz import total_lines_and_char
from extra import get_continuous_chunks



def preprocess_data(sen):
    sentences=[]
    sen=[el.strip() for el in sen.split("\n") if len(el) > 0]
    for sentence in sen:
        sentence=nltk.sent_tokenize(sentence)
        for a in sentence:
            sentences.append(a)
    #print(sentences)
    sen=[nltk.word_tokenize(sent) for sent in sentences]
    #print(sentences)
    sen=[nltk.pos_tag(sent) for sent in sen]
    return sen
#d=preprocess_data(document)
#print(d)



def extract_phone_numbers(string):
    r = re.compile(r'(\d{3}[-\.\s]??\d{3}[-\.\s]??\d{4}|\(\d{3}\)\s*\d{3}[-\.\s]??\d{4}|\d{3}[-\.\s]??\d{4})')
    phone_numbers = r.findall(string)
    phone_number=[f for f in phone_numbers if len(f)==10]
    return [re.sub(r'\D', '', number) for number in phone_numbers]
#numbers=extract_phone_numbers(document)
#print(numbers)



def extract_email_addresses(string):
    r = re.compile(r'[\w\.-]+@[\w\.-]+')
    return r.findall(string)
#emails=extract_email_addresses(document)
#print(emails)


# To get names ( still need some changes )

def get_names(string):
    data=preprocess_data(string)
    name=[]
    name2=[]
    try:
        for i in data:
            #print(i)
            chunkGram = r'Chunk: {<NNP><NNP><NNP>*}'
            chunkParser = nltk.RegexpParser(chunkGram)
            chunked = chunkParser.parse(i)
            #chunked.draw() 
            #print(chunked)
            # How to get access to this data, actually now the data is like a tree and chunked and no-chunked data are subtrees of tree
            #for subtree in chunked.subtrees():
                #print(subtree)
            # If we want only chunked variables
            for subtree in chunked.subtrees(filter=lambda t: t.label() == 'Chunk'):
                #print(subtree)
                name.append(' '.join([s[0] for s in subtree]))
            
            
        # data2=[nltk.word_tokenize(sent) for sent in name]
        # data2=[nltk.pos_tag(sent) for sent in data2]
        # for i in data2:
        #     # This nltk.ne_chunk extrcts proper noun from data and classifies them to which category they belong either person or organization or facility etc.
        #     namedEnt = nltk.ne_chunk(i, binary=False) # False to get the classified entities otherwise all entities would considered as NE(common)
        #     #namedEnt.draw()
        #     # To get names from a document
        #     for subtree in namedEnt.subtrees(filter=lambda t: t.label() == 'PERSON'):
        #         #print(subtree)
        #         name2.append(' '.join([s[0] for s in subtree]))
        for j in name:
            n=get_continuous_chunks(j)
            if(len(n)!=0):
                name2.append(n)

    except Exception as e:
        print(str(e))
    return name2
#a=get_names(document)
#print(a)

# for extracting the linked.in profile
def extract_linkedin(string):
    r = re.compile(r'[\w:\/\.]*linkedin.com\/[\w\/]{3,8}[\w-]{1,}[\/]?')  #http://www.linkedin.com/in/abc   linkedin.com/in/priyanshu-gupta07b802173
    return r.findall(string)
#profile=extract_linkedin(document)
#print(profile)


# for finding the no. of text lines + total no. of text characters on each page

def textLines_plus_char(path):
    ext=path.rsplit('.', 1)[1].lower()
    if(ext=="pdf"):
        total=total_lines_and_char(path)
        return total
    else:
        return None





# To count no. of tables
def count_tables(path):
    ext=path.rsplit('.', 1)[1].lower()
    #print(ext)
    if(ext=="docx"):
        wordDoc = Document(path)
        count=0
        for table in wordDoc.tables:
            count+=1
            # To extract data from table contents
            #for row in table.rows:
                #for cell in row.cells:
                    #print(cell.text)
        
    else:    # for PDF
        df = tabula.read_pdf(path, pages = 'all', multiple_tables = True)
        #print(df)
        count=0
        for table in df:
            count+=1
    return count
#noOfTables=count_tables(path)
#print(noOfTables)




# TO find which font styles used and which font sizes used in a document
def find_font(path):
    ext=path.rsplit('.', 1)[1].lower()
    #print(ext)
    if(ext=="docx"):
        document = Document(path)
        fname=[]
        size=[]
        for para in document.paragraphs:
            for run in para.runs:
                #print(run.font.name)
                #print(run.text)
                if(not(run.font.name==None)):
                    if(not(run.font.name in fname)):
                        fname.append(run.font.name)
                if(not(run.font.size==None)):
                    f=run.font.size
                    if(not(f.pt in size)):
                        #print(f.pt)
                        size.append(f.pt)
    else:
        # If file is pdf
        fname=font_style(path)
    return fname
#font_name=find_font(path)
#print(font_name)
#print(font_size)




# To count images present in a document
def count_images(filename):
    ext=filename.rsplit('.', 1)[1].lower()
    #print(ext)
    if(ext=="docx"):
        # for images
        count=0
        doc = Document(filename)
        for s in doc.inline_shapes:
            #print (s.height.cm,s.width.cm,s._inline.graphic.graphicData.pic.nvPicPr.cNvPr.name)
            count+=1
    else:
        count=image_finder(filename)
    return count
#noOfImages=count_images(path)
#print(noOfImages)


def export_to_csv(path):
    #path=r"C:\Users\Priyanshu Gupta\Desktop\Resume Extractor\priyanshu.docx"
    ext=path.rsplit('.', 1)[1].lower()
    if(ext=='docx'):
        document=docx2txt.process(path)
    else:
        document=convert_pdf_to_txt(path)
    #print(document)
    df=pd.DataFrame(columns=["Name","Email_Id","Phone_Nos","Linkedin_Profile","Textline+Totalchar on each Page","Font_Style","Font_Size","Table_Count","Image_Count"])
    df.loc[0]=[None,None,None,None,None,None,None,None,None]
    a=df.values
    name=get_names(document)
    # for loop
    if(len(name)>0):
        a[0][0]=name[0][0]
    emails=extract_email_addresses(document)
    if(len(emails)>0):
        a[0][1]=emails[0]
    numbers=extract_phone_numbers(document)
    if(len(numbers)>0):
        a[0][2]=numbers[0]
    profile=extract_linkedin(document)
    if(len(profile)>0):
        a[0][3]=profile[0]
    noOfTables=count_tables(path)
    a[0][7]=noOfTables
    font_name=find_font(path)
    a[0][5]=font_name
    noOfImages=count_images(path)
    a[0][8]=noOfImages
    total=textLines_plus_char(path)
    a[0][4]=total
    #print(df)
    csv_file_name=r"C:\Users\Priyanshu Gupta\Desktop\Resume Extractor"+"\\"+a[0][0]+" resume.csv"
    print(csv_file_name)
    csv_file=df.to_csv(csv_file_name,index=False)
    return(csv_file_name)
#export_to_csv(r"C:\Users\Priyanshu Gupta\Desktop\Resume Extractor\priyanshu.docx")
print(dir(nltk))





