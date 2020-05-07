from docx import Document
from docx.enum.text import WD_BREAK
fn=r"C:\Users\Priyanshu Gupta\Desktop\Resume Extractor\experiment4.docx"
document = Document(fn)
# pn=1    
# import re
# for p in document.paragraphs:
#     r=re.match('Chapter \d+',p.text)
#     if r:
#         print(r.group(),pn)
#     for run in p.runs:
#         if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
#             pn+=1
#             print(pn)
#             print('!!','='*50,pn)


for paragraph in document.paragraphs:
    sentences = paragraph.text.split('. ') 
    for sentence in sentences:
        words=sentence.split(' ')
        for word in words:
            
            
            for run in paragraph.runs:
                if run.page_break== WD_BREAK:
                    current_page_number +=1
            replace_counter += 1
            # write to a report what paragraph and what page
            write_report(error, correction, sentence, current_page_number )  
            # for that I need to know a page break    